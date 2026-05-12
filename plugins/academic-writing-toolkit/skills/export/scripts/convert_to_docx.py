#!/usr/bin/env python3
"""Convert thesis chapters and reading notes from Markdown to Word (.docx).

Usage:
    python convert_to_docx.py --base-dir /path/to/project --output-dir /path/to/output
    python convert_to_docx.py --base-dir . --output-dir ./export --scope chapters --lang-filter en-only

Conversion priority:
    1. pypandoc (if available)
    2. python-docx + markdown (fallback)
"""
from __future__ import annotations

import argparse
import os
import re
import sys
import zipfile
from datetime import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Conversion backends
# ---------------------------------------------------------------------------

USE_PANDOC = False

try:
    import pypandoc
    # Smoke probe: get_pandoc_version() shells out to `pandoc --version`
    # internally and raises if the binary is missing or broken
    # (Snap isolation, missing libs, broken wrapper).
    pypandoc.get_pandoc_version()
    USE_PANDOC = True
except (ImportError, OSError, RuntimeError):
    pass

# Always try to import python-docx — needed both as primary path when
# USE_PANDOC=False, and as runtime fallback when pypandoc fails on a
# specific file (see convert_file / create_cover try/except below).
HAS_DOCX = False
try:
    import markdown
    from docx import Document
    from docx.shared import Inches, Pt
    HAS_DOCX = True
except ImportError:
    pass

BACKEND_ERROR = (
    "Error: No conversion backend available.\n"
    "  - pypandoc is unavailable (not installed, or `pandoc` binary missing/broken on PATH).\n"
    "  - python-docx + markdown fallback is also missing.\n"
    "Install one of:\n"
    "  pip install pypandoc       (also requires `pandoc` binary on PATH)\n"
    "  pip install python-docx markdown"
)


def ensure_conversion_backend() -> None:
    """Exit with a clear message when no Markdown-to-docx backend is available."""
    if not USE_PANDOC and not HAS_DOCX:
        sys.exit(BACKEND_ERROR)


# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------

CJK_PATTERN = re.compile(r"[\u4e00-\u9fff]")


def cjk_ratio(text: str) -> float:
    """Return the proportion of CJK characters in *text*."""
    if not text:
        return 0.0
    cjk_count = len(CJK_PATTERN.findall(text))
    # Count only non-whitespace characters for the denominator
    total = len(re.sub(r"\s", "", text))
    if total == 0:
        return 0.0
    return cjk_count / total


def should_skip(text: str, lang_filter: str) -> bool:
    """Return True if the file should be skipped under *lang_filter*."""
    if lang_filter != "en-only":
        return False
    return cjk_ratio(text) > 0.01


# ---------------------------------------------------------------------------
# File conversion
# ---------------------------------------------------------------------------


def convert_file(src: Path, dst: Path) -> None:
    """Convert a single Markdown file to .docx."""
    dst.parent.mkdir(parents=True, exist_ok=True)

    if USE_PANDOC:
        try:
            pypandoc.convert_file(
                str(src),
                "docx",
                outputfile=str(dst),
                extra_args=["--wrap=none"],
            )
            return
        except (OSError, RuntimeError) as exc:
            print(f"  Warning: pandoc failed on {src.name}: {exc}; falling back to python-docx")
    _convert_with_docx(src, dst)


def _convert_with_docx(src: Path, dst: Path) -> None:
    """Fallback converter using python-docx + markdown."""
    if not HAS_DOCX:
        raise RuntimeError(
            "python-docx fallback not available; install python-docx + markdown"
        )
    md_text = src.read_text(encoding="utf-8")
    html = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "footnotes", "toc"],
    )

    doc = Document()

    # Apply a minimal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Simple HTML-to-docx: split by block tags and add paragraphs.
    # This is intentionally simple -- for full fidelity, use pandoc.
    _html_blocks_to_docx(doc, html)
    doc.save(str(dst))


def _html_blocks_to_docx(doc, html: str) -> None:
    """Minimal HTML block parser for python-docx fallback."""
    # Strip tags for a basic conversion; heading detection via markdown markers
    # We re-parse from a simple approach: line-by-line
    from html.parser import HTMLParser

    class _Parser(HTMLParser):
        def __init__(self, document):
            super().__init__()
            self.doc = document
            self.current_text = ""
            self.in_heading = 0  # heading level, 0 = not in heading
            self.in_table = False
            self.table_rows = []
            self.current_row = []
            self.in_cell = False

        def handle_starttag(self, tag, attrs):
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self.in_heading = int(tag[1])
                self.current_text = ""
            elif tag == "table":
                self.in_table = True
                self.table_rows = []
            elif tag == "tr":
                self.current_row = []
            elif tag in ("td", "th"):
                self.in_cell = True
                self.current_text = ""
            elif tag == "p":
                self.current_text = ""
            elif tag == "blockquote":
                self.current_text = ""
            elif tag == "br":
                self.current_text += "\n"

        def handle_endtag(self, tag):
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = min(self.in_heading, 4)  # docx supports heading 1-4 well
                self.doc.add_heading(self.current_text.strip(), level=level)
                self.in_heading = 0
                self.current_text = ""
            elif tag == "p":
                text = self.current_text.strip()
                if text:
                    self.doc.add_paragraph(text)
                self.current_text = ""
            elif tag == "blockquote":
                text = self.current_text.strip()
                if text:
                    p = self.doc.add_paragraph(text)
                    p.style = self.doc.styles["Quote"] if "Quote" in [
                        s.name for s in self.doc.styles
                    ] else self.doc.styles["Normal"]
                self.current_text = ""
            elif tag in ("td", "th"):
                self.in_cell = False
                self.current_row.append(self.current_text.strip())
                self.current_text = ""
            elif tag == "tr":
                self.table_rows.append(self.current_row)
            elif tag == "table":
                self.in_table = False
                if self.table_rows:
                    cols = max(len(r) for r in self.table_rows)
                    table = self.doc.add_table(
                        rows=len(self.table_rows), cols=cols
                    )
                    table.style = "Table Grid"
                    for i, row_data in enumerate(self.table_rows):
                        for j, cell_text in enumerate(row_data):
                            if j < cols:
                                table.rows[i].cells[j].text = cell_text

        def handle_data(self, data):
            if self.in_heading or self.in_cell:
                self.current_text += data
            else:
                self.current_text += data

    parser = _Parser(doc)
    parser.feed(html)


# ---------------------------------------------------------------------------
# Batch operations
# ---------------------------------------------------------------------------


def convert_chapters(base: Path, out: Path) -> tuple[int, int]:
    """Convert all chapter files. Returns (converted, skipped)."""
    chapters_dir = base / "chapters"
    if not chapters_dir.is_dir():
        print(f"Warning: {chapters_dir} not found, skipping chapters.")
        return 0, 0
    return _convert_dir(chapters_dir, out / "chapters", "ch*.md")


def convert_notes(base: Path, out: Path, lang_filter: str) -> tuple[int, int]:
    """Convert reading notes. Returns (converted, skipped)."""
    notes_dir = base / "literature" / "reading_notes"
    if not notes_dir.is_dir():
        print(f"Warning: {notes_dir} not found, skipping notes.")
        return 0, 0
    return _convert_dir(notes_dir, out / "notes", "*_NOTES.md", lang_filter)


def _convert_dir(
    src_dir: Path, dst_dir: Path, pattern: str, lang_filter: str = "all"
) -> tuple[int, int]:
    """Convert all matching files in a directory. Returns (converted, skipped)."""
    converted = 0
    skipped = 0
    for md_file in sorted(src_dir.glob(pattern)):
        text = md_file.read_text(encoding="utf-8")
        if should_skip(text, lang_filter):
            print(f"  Skipped (language filter): {md_file.name}")
            skipped += 1
            continue
        dst = dst_dir / (md_file.stem + ".docx")
        print(f"  Converting: {md_file.name} -> {dst.name}")
        convert_file(md_file, dst)
        converted += 1
    return converted, skipped


def create_cover(metadata: dict, out: Path) -> Path:
    """Create a simple cover page .docx with thesis metadata."""
    dst = out / "00_cover.docx"
    dst.parent.mkdir(parents=True, exist_ok=True)

    if USE_PANDOC:
        try:
            cover_md = f"# {metadata.get('title', 'Thesis')}\n\n"
            cover_md += f"**Author**: {metadata.get('author', 'Unknown')}\n\n"
            cover_md += f"**Date**: {metadata.get('date', datetime.now().strftime('%Y-%m-%d'))}\n\n"
            cover_md += f"**Institution**: {metadata.get('institution', '')}\n\n"
            if metadata.get("abstract"):
                cover_md += f"## Abstract\n\n{metadata['abstract']}\n"
            tmp = out / "_cover_tmp.md"
            tmp.write_text(cover_md, encoding="utf-8")
            pypandoc.convert_file(str(tmp), "docx", outputfile=str(dst))
            tmp.unlink()
            return dst
        except (OSError, RuntimeError) as exc:
            print(f"  Warning: pandoc failed on cover: {exc}; falling back to python-docx")
            if tmp.exists():
                tmp.unlink()

    # Fallback path
    if not HAS_DOCX:
        raise RuntimeError("python-docx fallback not available")
    doc = Document()
    doc.add_heading(metadata.get("title", "Thesis"), level=0)
    doc.add_paragraph(f"Author: {metadata.get('author', 'Unknown')}")
    doc.add_paragraph(
        f"Date: {metadata.get('date', datetime.now().strftime('%Y-%m-%d'))}"
    )
    doc.add_paragraph(f"Institution: {metadata.get('institution', '')}")
    if metadata.get("abstract"):
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(metadata["abstract"])
    doc.save(str(dst))

    return dst


def package_zip(out_dir: Path, zip_path: Path) -> None:
    """Create a ZIP archive of the output directory with UTF-8 filenames."""
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _dirs, files in os.walk(out_dir):
            for fname in sorted(files):
                if fname.endswith(".zip"):
                    continue
                filepath = Path(root) / fname
                # Use forward slashes and English folder names
                arcname = filepath.relative_to(out_dir).as_posix()
                # Set UTF-8 flag (bit 11) via ZipInfo
                info = zipfile.ZipInfo(arcname)
                info.flag_bits |= 0x800  # bit 11 = UTF-8
                info.compress_type = zipfile.ZIP_DEFLATED
                with open(filepath, "rb") as f:
                    zf.writestr(info, f.read())
    print(f"  ZIP created: {zip_path} ({zip_path.stat().st_size / 1024:.1f} KB)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Convert thesis Markdown files to Word (.docx)"
    )
    parser.add_argument(
        "--base-dir",
        type=Path,
        required=True,
        help="Project root directory containing chapters/ and literature/",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        required=True,
        help="Output directory for converted files",
    )
    parser.add_argument(
        "--scope",
        choices=["chapters", "notes", "all"],
        default="all",
        help="What to convert (default: all)",
    )
    parser.add_argument(
        "--lang-filter",
        choices=["en-only", "all"],
        default="all",
        help="Language filter for notes (default: all)",
    )
    args = parser.parse_args()
    ensure_conversion_backend()

    base = args.base_dir.resolve()
    out = args.output_dir.resolve()

    print(f"Base directory: {base}")
    print(f"Output directory: {out}")
    print(f"Scope: {args.scope}")
    print(f"Language filter: {args.lang_filter}")
    print(f"Conversion method: {'pypandoc (pandoc)' if USE_PANDOC else 'python-docx + markdown (fallback)'}")
    print()

    total_converted = 0
    total_skipped = 0

    if args.scope in ("chapters", "all"):
        print("--- Chapters ---")
        c, s = convert_chapters(base, out)
        total_converted += c
        total_skipped += s
        print(f"  Chapters: {c} converted, {s} skipped\n")

    if args.scope in ("notes", "all"):
        print("--- Reading Notes ---")
        c, s = convert_notes(base, out, args.lang_filter)
        total_converted += c
        total_skipped += s
        print(f"  Notes: {c} converted, {s} skipped\n")

    if total_converted > 0:
        # Package into ZIP
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_path = out / f"thesis_export_{timestamp}.zip"
        print("--- Packaging ---")
        package_zip(out, zip_path)

    print(f"\nDone. {total_converted} files converted, {total_skipped} skipped.")


if __name__ == "__main__":
    main()
